#Contains code for automatically generating a decently formatted docx document with selected chemicals
#The functions in here are not meant to edit/add new chemicals to the JSON, only read from it

from docx import Document       #main functions to allow docx creation
from docx.shared import Cm      #subfunction for length - here Cm imported allows for specifying document lengths in centimeters - Inches for inches

#Used object structure as to have an internally 'global' document (self.document) that can be updated rather than constantly passing and returning a document file
class Experiment_doc:
    def __init__(self, title, chemicals_list, sort_type=None):
        self.title = title
        self.document = Document()
        self.list = chemicals_list      #When a document is generated, the user selects a number of chemicals to include which are in this list by key
        self.sort = sort_type           #Option to group chemicals by inclusion order, manual numbering by the user or by grouping together by type

        if self.sort == 'Type':         #If the user wants the chemicals to appear in groups based on type (e.g, mineral acids), this creates a dict of all the types for later iteration
            types = {}                  
            for chemical in self.list:
                types[self.list[chemical]['Type']] = True
            self.types = types
         
    def generate_doc(self):
        #if header in -local directory goes here-:      the docx functions will overwrite existing documents if the title matches -add a check to avoid
        self.document.add_heading(self.title, level=0)
        for chemical in self.list:
            self.document.add_heading(chemical, level=2)
            self.add_chemical(chemical)
            self.document.add_paragraph('----------------------------------------------------------------------------------------------------------------------  ') #Add a divider between entries
                                        #Needs a way to add page breaks for proper formatting
        self.document.save(self.title)

    #The paragraph additions are currently not superbly well formated - fix
    def add_chemical(self, chemical):
        paragraph = self.document.add_paragraph()
        entry = self.list[chemical]
        #tab_stop = tab_stop.add_tab_stop(Cm(8.5))      function to add a tab stop to the document
        count = 0   #Basic, kinda crap method to count each data point added so that linebreaks can be inserted properly
        for key in entry:               #Iterate over the related chemical and add all the different data fields to the document
            paragraph.add_run(key)
            paragraph.add_run(': ')
            paragraph.add_run(entry[key])
            count += 1
            if count % 2 == 0:
                paragraph.add_run().add_break()     #add_break default adds a linebreak, can set to other types, accessed under add_run
            else:
                paragraph.add_run().add_tab()       #adds a tab character from under add_run - add_tab_stop should be useable here
                paragraph.add_run().add_tab()       #Or find a different way to move the next section of text to a specific point - double tab results in uneven text